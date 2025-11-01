import re
from docx import Document
import traceback
import os
import uuid
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from docxtpl import DocxTemplate

# --- App Setup ---
app = Flask(__name__)
# Allow requests from our React frontend (which will run on http://localhost:3000)
CORS(app, resources={r"/*": {"origins": ["http://localhost:3000", "https://docu-filler-qk4hne7rv-soham-katkars-projects.vercel.app"]}})

# --- Constants ---
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# --- Helper Function: Extract Placeholders (with DEBUGGING) ---
# --- Helper Function: Extract Placeholders (python-docx + REGEX METHOD) ---
def get_placeholders_from_template(template_path):
    """
    Uses python-docx directly to read text and Regex to find placeholders.
    This bypasses the broken DocxTemplate() initialization.
    """
    try:
        print("--- DEBUG: Opening with python-docx directly ---")
        doc = Document(template_path)
        full_text = []
        
        # Read all text from all paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # You could also add text from tables if your templates have them
        # for table in doc.tables:
        #     for row in table.rows:
        #         for cell in row.cells:
        #             full_text.append(cell.text)
        
        xml_content = "\n".join(full_text)
        print("--- DEBUG: Successfully read all text. ---")

        # Use a regular expression to find all Jinja2-style variables
        regex = r'\{\{\s*([a-zA-Z0-9_]+)\s*\}\}'
        
        print("--- DEBUG: Running regex... ---")
        matches = re.findall(regex, xml_content)
        
        # Get a unique, sorted list of keys
        keys = sorted(list(set(matches)))
        
        if not keys:
            print("--- DEBUG: No placeholder keys found. ---")
            return []
            
        print(f"--- DEBUG: Found keys via regex: {keys} ---")
        
        placeholders = []
        for key in keys:
            label = key.replace('_', ' ').title()
            placeholders.append({
                "key": key,
                "label": label,
                "prompt": f"Please enter the {label}:",
                "value": ""
            })
        return placeholders
        
    except Exception as e:
        print(f"--- DEBUG: An exception occurred in python-docx REGEX method ---")
        traceback.print_exc() 
        print(f"Error parsing template: {e}")
        return None

# --- API Endpoint 1: Upload & Parse ---
@app.route('/upload', methods=['POST'])
def upload_template():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    if file and file.filename.endswith('.docx'):
        # 1. Create a unique session ID for this file
        file_id = str(uuid.uuid4())
        session_folder = os.path.join(app.config['UPLOAD_FOLDER'], file_id)
        os.makedirs(session_folder)
        
        # 2. Save the original template
        template_path = os.path.join(session_folder, "template.docx")
        file.save(template_path)
        
        # 3. Parse placeholders
        placeholders = get_placeholders_from_template(template_path)
        
        if placeholders is None:
            return jsonify({"error": "Could not parse document. Is it a valid .docx template?"}), 500

        # 4. Return the file ID and the list of questions
        return jsonify({
            "fileId": file_id,
            "placeholders": placeholders
        }), 200
    
    return jsonify({"error": "Invalid file type. Please upload a .docx file."}), 400

# --- API Endpoint 2: Generate & Fill ---
@app.route('/generate', methods=['POST'])
def generate_document():
    data = request.get_json()
    file_id = data.get('fileId')
    answers = data.get('answers') # This is our {'COMPANY_NAME': 'Acme Inc.', ...}
    
    if not file_id or not answers:
        return jsonify({"error": "Missing fileId or answers"}), 400
    
    template_path = os.path.join(app.config['UPLOAD_FOLDER'], file_id, "template.docx")
    output_path = os.path.join(app.config['UPLOAD_FOLDER'], file_id, "filled_document.docx")
    
    if not os.path.exists(template_path):
        return jsonify({"error": "File session not found. Please upload again."}), 404

    # 1. Load the template
    doc = DocxTemplate(template_path)
    
    # 2. Render the answers into the template
    doc.render(answers)
    
    # 3. Save the new filled document
    doc.save(output_path)
    
    # 4. Return a URL to download the new file
    download_url = f"http://127.0.0.1:5000/download/{file_id}"
    return jsonify({"downloadUrl": download_url}), 200

# --- API Endpoint 3: Download File ---
@app.route('/download/<file_id>', methods=['GET'])
def download_file(file_id):
    directory = os.path.join(app.config['UPLOAD_FOLDER'], file_id)
    return send_from_directory(directory, "filled_document.docx", as_attachment=True)

# --- Run the App ---
if __name__ == '__main__':
    app.run(debug=True, port=5000)